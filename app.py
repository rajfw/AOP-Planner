import os
import json
import logging
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_from_directory, send_file, redirect, url_for, session
from flask_cors import CORS
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from authlib.integrations.flask_client import OAuth
from werkzeug.utils import secure_filename
import tempfile
from pathlib import Path
import markdown
from dotenv import load_dotenv
import io
import uuid
import shutil

# Document parsing libraries
import docx
import PyPDF2
import fitz  # PyMuPDF

# LLM integration - Updated for OpenAI v1.x
try:
    from openai import OpenAI
    import tenacity
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("OpenAI package not available. Install with: pip install openai")

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)

# Configuration
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['ALLOWED_EXTENSIONS'] = {'txt', 'pdf', 'doc', 'docx', 'md'}

# Create uploads and roadmap directories
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROADMAP_DATA_DIR = os.path.join(BASE_DIR, 'data')
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'uploads')

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
ROADMAP_UPLOAD_DIR = os.path.join(app.config['UPLOAD_FOLDER'], 'roadmap_attachments')
os.makedirs(ROADMAP_DATA_DIR, exist_ok=True)
os.makedirs(ROADMAP_UPLOAD_DIR, exist_ok=True)

ROADMAP_FILE = os.path.join(ROADMAP_DATA_DIR, 'roadmaps.json')
if not os.path.exists(ROADMAP_FILE):
    with open(ROADMAP_FILE, 'w') as f:
        json.dump([], f)

def get_roadmaps():
    try:
        with open(ROADMAP_FILE, 'r') as f:
            return json.load(f)
    except:
        return []

def save_roadmaps(roadmaps):
    with open(ROADMAP_FILE, 'w') as f:
        json.dump(roadmaps, f, indent=2)

# User Database Initialization
USERS_FILE = os.path.join(ROADMAP_DATA_DIR, 'users.json')

# Auth Setup
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'fw-aop-secret-key-2025')
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=os.getenv('GOOGLE_CLIENT_ID'),
    client_secret=os.getenv('GOOGLE_CLIENT_SECRET'),
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={'scope': 'openid email profile'},
)

class User(UserMixin):
    def __init__(self, user_data):
        self.id = user_data['id']
        self.username = user_data.get('username')
        self.name = user_data.get('name')
        self.email = user_data.get('email')
        self.role = user_data.get('role', 'requester')
        self.avatar = user_data.get('avatar')

def get_users():
    print(f"DEBUG: get_users called. Path: {USERS_FILE}")
    if os.path.exists(USERS_FILE):
        try:
            with open(USERS_FILE, 'r') as f:
                data = json.load(f)
                print(f"DEBUG: Loaded {len(data)} users: {[u.get('username') for u in data]}")
                return data
        except Exception as e:
            print(f"DEBUG: Error loading users.json: {e}")
            return []
    else:
        print(f"DEBUG: Users file NOT found at {USERS_FILE}")
        # data dir check
        print(f"DEBUG: Data dir exists? {os.path.exists(ROADMAP_DATA_DIR)}")
        print(f"DEBUG: CWD: {os.getcwd()}")
    return []

@login_manager.user_loader
def load_user(user_id):
    users = get_users()
    user_data = next((u for u in users if u['id'] == user_id), None)
    if user_data:
        return User(user_data)
    return None

from functools import wraps
from flask import abort

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != 'admin':
            abort(403)
        return f(*args, **kwargs)
    return decorated_function

# Form Configuration Management
FORM_CONFIG_FILE = os.path.join(ROADMAP_DATA_DIR, 'roadmap_form.json')

def get_form_config():
    if os.path.exists(FORM_CONFIG_FILE):
        with open(FORM_CONFIG_FILE, 'r') as f:
            return json.load(f)
    return []

def save_form_config(config):
    with open(FORM_CONFIG_FILE, 'w') as f:
        json.dump(config, f, indent=2)

# Initialize OpenAI client - FIXED VERSION
if OPENAI_AVAILABLE:
    try:
        # Initialize with proper configuration for v1.x
        import openai
        # Initialize with proper configuration for v1.x
        client = OpenAI(
            api_key="eyJhbGciOiJIUzI1NiJ9.eyJ1c2VyIjp7Im5hbWUiOiJSYWplc3dhciBQIFMiLCJlbWFpbCI6InJhamVzd2FyLnN1YnJhbWFuaUBmcmVzaHdvcmtzLmNvbSIsImltYWdlIjoiaHR0cHM6Ly9saDMuZ29vZ2xldXNlcmNvbnRlbnQuY29tL2EvQUNnOG9jS0FBUkNWSktyMjhxbU0xRTdnUE1fSlhPcDU4MEZHM2prNThMYzQ1SVB6eVFqN0lxWF89czk2LWMifSwianRpIjoiUGNVU0xxOFNrR3lYdmF1aFlEQTdsIiwiaWF0IjoxNzcwNzg2NTYwLCJleHAiOjE3NzEzOTEzNjB9.JglVKNUeldw7thE2swT0jiXKkf2M3DNUCZZ0WAIAWOg",
            base_url="https://cloudverse.freshworkscorp.com/api/v1"
        )
        print("OpenAI client initialized successfully")
    except Exception as e:
        print(f"Error initializing OpenAI client: {e}")
        client = None
else:
    client = None
    print("OpenAI package not available. LLM features will be disabled.")

# Configure logging
# Configure logging
logging.basicConfig(
    level=logging.INFO,
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def parse_document(filepath, filename):
    """Parse different document formats and extract text."""
    ext = filename.rsplit('.', 1)[1].lower()
    text = ""
    
    try:
        if ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
                
        elif ext == 'pdf':
            # Method 1: Try PyPDF2
            try:
                with open(filepath, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except:
                # Method 2: Try PyMuPDF
                try:
                    doc = fitz.open(filepath)
                    for page in doc:
                        text += page.get_text() + "\n"
                except Exception as e:
                    logger.error(f"PDF parsing error: {e}")
                    text = f"Error parsing PDF: {str(e)}"
                    
        elif ext in ['doc', 'docx']:
            try:
                doc = docx.Document(filepath)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
            except Exception as e:
                logger.error(f"DOCX parsing error: {e}")
                text = f"Error parsing DOCX: {str(e)}"
                
        elif ext == 'md':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
                
    except Exception as e:
        logger.error(f"Error parsing document {filename}: {e}")
        text = f"Error parsing document: {str(e)}"
    
    return {
        'filename': filename,
        'content': text,
        'word_count': len(text.split()),
        'char_count': len(text),
        'parsed_at': datetime.now().isoformat()
    }

def parse_prd_structure(text):
    """Parse PRD and extract structured information."""
    sections = {
        'overview': '',
        'objectives': '',
        'scope': '',
        'features': '',
        'user_stories': '',
        'requirements': '',
        'success_metrics': '',
        'timeline': '',
        'risks': ''
    }
    
    # Simple section detection based on common headers
    lines = text.split('\n')
    current_section = None
    
    for line in lines:
        line_lower = line.strip().lower()
        
        # Detect section headers
        if any(keyword in line_lower for keyword in ['overview', 'introduction', 'background']):
            current_section = 'overview'
        elif any(keyword in line_lower for keyword in ['objectives', 'goals', 'purpose']):
            current_section = 'objectives'
        elif any(keyword in line_lower for keyword in ['scope', 'in scope', 'out of scope']):
            current_section = 'scope'
        elif any(keyword in line_lower for keyword in ['features', 'functionality']):
            current_section = 'features'
        elif any(keyword in line_lower for keyword in ['user stories', 'user personas', 'user journey']):
            current_section = 'user_stories'
        elif any(keyword in line_lower for keyword in ['requirements', 'functional requirements', 'non-functional']):
            current_section = 'requirements'
        elif any(keyword in line_lower for keyword in ['success metrics', 'kpis', 'metrics']):
            current_section = 'success_metrics'
        elif any(keyword in line_lower for keyword in ['timeline', 'milestones', 'schedule']):
            current_section = 'timeline'
        elif any(keyword in line_lower for keyword in ['risks', 'assumptions', 'constraints']):
            current_section = 'risks'
        elif line.strip() and current_section:
            # Add content to current section
            sections[current_section] += line + '\n'
    
    return sections

@tenacity.retry(
    stop=tenacity.stop_after_attempt(3),
    wait=tenacity.wait_exponential(multiplier=1, min=4, max=10)
)
def improve_prd_with_llm(prd_content, improvement_type="comprehensive"):
    """Use LLM to improve PRD content."""
    if not client:
        return {"error": "OpenAI API key not configured or client not available"}
    
    try:
        prompt_templates = {
            "comprehensive": """
            As a product management expert, improve this Product Requirements Document (PRD).
            
            Original PRD:
            {prd_content}
            
            Please provide:
            1. Enhanced executive summary
            2. Clearer objectives and success metrics
            3. More detailed user stories with acceptance criteria
            4. Better structured requirements (functional & non-functional)
            5. Improved risk assessment and mitigation strategies
            6. Additional sections if needed
            
            Return the improved PRD in markdown format with proper sections.
            """,
            
            "clarify": """
            Clarify and simplify this PRD. Make it more concise and easier to understand:
            
            {prd_content}
            
            Focus on:
            - Simplifying language
            - Removing ambiguity
            - Improving structure
            - Adding clear definitions
            """,
            
            "expand": """
            Expand this PRD with more details and depth:
            
            {prd_content}
            
            Add:
            - More detailed user personas
            - Additional edge cases
            - Technical considerations
            - Implementation details
            - Testing scenarios
            """
        }
        
        prompt = prompt_templates.get(
            improvement_type, 
            prompt_templates["comprehensive"]
        ).format(prd_content=prd_content)
        
        response = client.chat.completions.create(
            model="Azure-GPT-5-chat",  # Updated to user-provided model name
            messages=[
                {"role": "system", "content": "You are an expert product manager with experience at top tech companies."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        
        improved_content = response.choices[0].message.content
        return {
            "improved_content": improved_content,
            "model": "Azure-GPT-5-chat",
            "improvement_type": improvement_type,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"LLM improvement error: {e}")
        return {"error": f"Failed to improve PRD: {str(e)}"}

@app.route('/login')
def login():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    return render_template('login.html')

@app.route('/login/google')
def google_login():
    redirect_uri = url_for('google_callback', _external=True)
    return google.authorize_redirect(redirect_uri)

@app.route('/login/google/callback')
def google_callback():
    token = google.authorize_access_token()
    user_info = google.get('userinfo').json()
    
    users = get_users()
    user_data = next((u for u in users if u['email'] == user_info['email']), None)
    
    if not user_data:
        # Create new user for first-time Google login
        user_data = {
            "id": str(uuid.uuid4()),
            "username": user_info['email'],
            "name": user_info['name'],
            "email": user_info['email'],
            "role": "requester",
            "avatar": user_info.get('picture', f"https://ui-avatars.com/api/?name={user_info['name']}")
        }
        users.append(user_data)
        with open(USERS_FILE, 'w') as f:
            json.dump(users, f, indent=2)
            
    user = User(user_data)
    login_user(user)
    return redirect(url_for('home'))

@app.route('/api/login', methods=['POST'])
def api_login():
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    
    print(f"Login attempt - Username: '{username}'")
    
    users = get_users()
    
    # Debug finding user
    found_user = next((u for u in users if u['username'] == username), None)
    if found_user:
        print(f"User '{username}' exists. Checking password...")
        if found_user['password'] == password:
            print("Password match!")
        else:
            print(f"Password mismatch. Input length: {len(password)}, Stored length: {len(found_user['password'])}")
    else:
        print(f"User '{username}' does not exist in loaded users list.")
        
    user_data = next((u for u in users if u['username'] == username and u['password'] == password), None)
    
    if user_data:
        user = User(user_data)
        login_user(user)
        return jsonify({'success': True})
    
    return jsonify({'success': False, 'error': 'Invalid credentials'}), 401

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/admin')
@admin_required
def admin_panel():
    return render_template('admin.html')

@app.route('/api/form-config', methods=['GET'])
@login_required
def get_form_config_api():
    return jsonify(get_form_config())

@app.route('/api/form-config', methods=['POST'])
@admin_required
def update_form_config_api():
    try:
        config = request.json
        save_form_config(config)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/load-sample', methods=['POST'])
@admin_required
def load_sample_data():
    try:
        # Sample data provided by user (formatted as a list of dicts for local processing)
        sample_raw = [
            {"name": "Forms 2.0: Richer Fields, Seamless Migration, Unified Experience", "desc": "Introduce Rich Text Editor, Attachment, and Formula fields in all forms...", "bu": "EX", "type": "Major Feature", "priority": "P2", "q": "Q2"},
            {"name": "Automated Change Risk scoring and Assessment", "desc": "Introduce a native, scalable risk assessment framework within Change Management...", "bu": "EX", "type": "Major Feature", "priority": "P1", "q": "Q1"},
            {"name": "Service Health Dashboard powered by D42-ITOM Integration", "desc": "A more powerful and feature-rich Service Health Dashboard powered by Device42 and ITOM", "bu": "EX", "type": "Major Feature", "priority": "P1", "q": "Q1"},
            {"name": "New FD Omni for existing customers", "desc": "Migration readiness for customers from FD(older accounts), Classic Omni & CSS Omni...", "bu": "CX", "type": "Major Feature", "priority": "P1", "q": "Q2"},
            {"name": "Proactive messaging for Omnichannel Freshdesk", "desc": "Support for contact segmentation and proactive campaigns across channels...", "bu": "CX", "type": "Major Feature", "priority": "P2", "q": "Q3"},
            {"name": "Enhanced portal", "desc": "AI powered search on the portal, AI prompt portal customization, OOTB themes...", "bu": "CX", "type": "Major Feature", "priority": "P2", "q": "Q2"},
            {"name": "Enhanced Knowledge base", "desc": "AI powered KB insights and article generation, AI powered translation...", "bu": "CX", "type": "Major Feature", "priority": "P2", "q": "Q3"},
            {"name": "Multi-departments", "desc": "Improved RBAC and isolation for tickets, contacts, analytics and admin setup...", "bu": "CX", "type": "Major Feature", "priority": "P2", "q": "Q1"}
        ]

        roadmaps = get_roadmaps()
        
        for s in sample_raw:
            # Map BU
            bu_map = {"EX": "EX BU", "AI": "AI BU", "CX": "CX BU", "CE": "CE BU"}
            bu = bu_map.get(s['bu'], "AI BU")
            
            # Map Impact
            impact_map = {"P1": 5, "P2": 4, "P3": 3}
            impact = impact_map.get(s['priority'], 3)
            
            # Map Type
            f_type = "Hero Big Rock" if "Major" in s['type'] else "Big Rock"
            
            new_item = {
                "id": str(uuid.uuid4()),
                "title": s['name'],
                "description": s['desc'],
                "business_unit": bu,
                "target_year": "2026",
                "half_year": "H1" if s['q'] in ['Q1', 'Q2'] else "H2",
                "quarter": s['q'],
                "feature_type": f_type,
                "business_impact": impact,
                "dependencies": [],
                "attachments": [],
                "created_at": datetime.now().isoformat(),
                "created_by": current_user.name
            }
            roadmaps.append(new_item)
            
        save_roadmaps(roadmaps)
        return jsonify({'success': True, 'count': len(sample_raw)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/')
@login_required
def home():
    return render_template('home.html')

@app.route('/prd')
@login_required
def prd_tool():
    return render_template('prd_management.html')

@app.route('/roadmap')
@login_required
def roadmap_tool():
    return render_template('roadmap.html')

@app.route('/api/upload', methods=['POST'])
def upload_prd():
    """Handle PRD file upload and parsing."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        try:
            # Parse the document
            parsed_data = parse_document(filepath, filename)
            
            # Extract structure
            structure = parse_prd_structure(parsed_data['content'])
            
            # Save metadata
            metadata = {
                'original_filename': filename,
                'uploaded_at': datetime.now().isoformat(),
                'file_size': os.path.getsize(filepath),
                'parsed_data': parsed_data,
                'structure': structure
            }
            
            # Save metadata to file
            metadata_file = os.path.join(app.config['UPLOAD_FOLDER'], f"{filename}.meta.json")
            with open(metadata_file, 'w') as f:
                json.dump(metadata, f, indent=2)
            
            return jsonify({
                'success': True,
                'filename': filename,
                'parsed_data': parsed_data,
                'structure': structure,
                'metadata': metadata
            })
            
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500
    
    return jsonify({'error': 'File type not allowed'}), 400

@app.route('/api/improve', methods=['POST'])
@login_required
def improve_prd():
    """Improve PRD using LLM."""
    data = request.json
    prd_content = data.get('content', '')
    improvement_type = data.get('improvement_type', 'comprehensive')
    
    if not prd_content:
        return jsonify({'error': 'No content provided'}), 400
    
    try:
        result = improve_prd_with_llm(prd_content, improvement_type)
        
        if 'error' in result:
            return jsonify(result), 500
        
        return jsonify({
            'success': True,
            'improved_content': result['improved_content'],
            'metadata': {
                'model': result['model'],
                'improvement_type': result['improvement_type'],
                'timestamp': result['timestamp']
            }
        })
        
    except Exception as e:
        logger.error(f"Improvement error: {e}")
        return jsonify({'error': f'Improvement failed: {str(e)}'}), 500

@app.route('/api/save', methods=['POST'])
@login_required
def save_prd():
    """Save edited PRD."""
    data = request.json
    content = data.get('content', '')
    filename = data.get('filename', f'prd_edited_{datetime.now().strftime("%Y%m%d_%H%M%S")}.md')
    
    if not content:
        return jsonify({'error': 'No content to save'}), 400
    
    try:
        # Save the edited content
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write(content)
            
        # Parse for metadata consistency
        structure = parse_prd_structure(content)
        metadata = {
            'original_filename': filename,
            'uploaded_at': datetime.now().isoformat(),
            'file_size': len(content.encode('utf-8')),
            'parsed_data': {
                'content': content,
                'filename': filename
            },
            'structure': structure
        }
        
        # Save metadata to file
        metadata_file = save_path + '.meta.json'
        with open(metadata_file, 'w') as f:
            json.dump(metadata, f, indent=2)
        
        return jsonify({
            'success': True,
            'filename': filename,
            'saved_at': datetime.now().isoformat(),
            'filepath': save_path,
            'metadata': metadata
        })
        
    except Exception as e:
        logger.error(f"Save error: {e}")
        return jsonify({'error': f'Save failed: {str(e)}'}), 500

def get_rice_analysis(content):
    """Extract RICE scores from PRD content using AI."""
    if not client:
        return None
        
    try:
        response = client.chat.completions.create(
            model="Azure-GPT-5-chat",
            messages=[
                {"role": "system", "content": "You are a product management consultant. Analyze the provided PRD and extract RICE (Reach, Impact, Confidence, Effort) scores.\n\nReach: Number of users/customers affected (e.g., 500, 10000).\nImpact: Value to the user (0.25 to 3.0: 3=Massive, 2=High, 1=Medium, 0.5=Low, 0.25=Minimal).\nConfidence: Percentage certainty (50% to 100%).\nEffort: Person-months (e.g., 0.5, 3.0).\n\nReturn ONLY a JSON object with keys: reach, impact, confidence, effort, total_score, and verdict (a 1-sentence strategic summary)."},
                {"role": "user", "content": f"Analyze this PRD for RICE scores:\n\n{content}"}
            ],
            response_format={ "type": "json_object" }
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        logger.error(f"RICE extraction error: {e}")
        return None

@app.route('/api/analyze', methods=['POST'])
@login_required
def analyze_prd():
    """Analyze PRD for completeness and RICE prioritization."""
    data = request.json
    prd_content = data.get('content', '')
    
    if not prd_content:
        return jsonify({'error': 'No content provided'}), 400
    
    try:
        # 1. Structure Analysis
        structure = parse_prd_structure(prd_content)
        
        # 2. Completeness Check
        missing_sections = []
        for section, content in structure.items():
            if not content.strip():
                missing_sections.append(section.replace('_', ' ').title())
        
        total_sections = len(structure)
        filled_sections = sum(1 for content in structure.values() if content.strip())
        completeness_score = (filled_sections / total_sections) * 100
        
        # 3. RICE Analysis (New!)
        rice_data = get_rice_analysis(prd_content)
        
        # 4. Final Payload
        return jsonify({
            'success': True,
            'analysis': {
                'completeness_score': round(completeness_score, 2),
                'missing_sections': missing_sections,
                'word_count': len(prd_content.split()),
                'section_counts': {k: len(v.split()) for k, v in structure.items()},
                'rice': rice_data,
                'recommendations': [
                    "Add clear success metrics" if not structure['success_metrics'] else "",
                    "Define user personas" if not structure['user_stories'] else "",
                    "Include timeline estimates" if not structure['timeline'] else ""
                ]
            }
        })
        
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

@app.route('/api/prd/chat', methods=['POST'])
@login_required
def prd_assistant_chat():
    """Context-aware AI assistant for PRD refinement."""
    if not client:
        return jsonify({'success': False, 'error': 'AI features are currently disabled.'})
        
    data = request.json
    messages = data.get('messages', []) # Chat history
    prd_context = data.get('prd_context', '') # Current editor content
    
    if not messages:
        return jsonify({'success': False, 'error': 'No messages provided'})

    try:
        # Construct system prompt with PRD context
        system_message = {
            "role": "system",
            "content": (
                "You are an expert Product Manager assistant helping to refine a Product Requirements Document (PRD).\n"
                "Your goal is to provide specific, actionable advice or content for the PRD.\n"
                "When asked to write or improve a section, keep it consistent with the existing tone.\n"
                "CURRENT PRD CONTEXT:\n"
                f"\"\"\"\n{prd_context}\n\"\"\"\n\n"
                "If the user asks for a new section, provide it in clear Markdown format."
            )
        }
        
        # Prepare full message list for LLM
        llm_messages = [system_message] + messages
        
        response = client.chat.completions.create(
            model=app.config.get('OPENAI_MODEL', "Azure-GPT-5-chat"),
            messages=llm_messages,
            temperature=0.7
        )
        
        reply = response.choices[0].message.content
        return jsonify({'success': True, 'reply': reply})
        
    except Exception as e:
        logger.error(f"Chat error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/list', methods=['GET'])
def list_prds():
    """List all uploaded PRDs."""
    try:
        prds = []
        for filename in os.listdir(app.config['UPLOAD_FOLDER']):
            if not filename.endswith('.meta.json'):
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                metadata_file = filepath + '.meta.json'
                
                if os.path.exists(metadata_file):
                    with open(metadata_file, 'r') as f:
                        metadata = json.load(f)
                    prds.append({
                        'filename': filename,
                        'metadata': metadata
                    })
        
        return jsonify({
            'success': True,
            'prds': prds,
            'count': len(prds)
        })
        
    except Exception as e:
        logger.error(f"List error: {e}")
        return jsonify({'error': f'Failed to list PRDs: {str(e)}'}), 500

@app.route('/api/delete/<filename>', methods=['DELETE'])
@login_required
def delete_prd(filename):
    """Delete a PRD and its metadata."""
    try:
        # Security check to ensure filename is safe
        if not allowed_file(filename) and not filename.endswith('.meta.json'):
             return jsonify({'error': 'Invalid filename'}), 400
             
        filename = secure_filename(filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        metadata_file = filepath + '.meta.json'
        
        # Check if file exists
        if not os.path.exists(filepath):
             return jsonify({'error': 'File not found'}), 404
             
        # Delete file
        os.remove(filepath)
        
        # Delete metadata if exists
        if os.path.exists(metadata_file):
            os.remove(metadata_file)
            
        return jsonify({
            'success': True,
            'message': f'Successfully deleted {filename}'
        })
        
    except Exception as e:
        logger.error(f"Delete error: {e}")
        return jsonify({'error': f'Failed to delete file: {str(e)}'}), 500

@app.route('/api/export', methods=['POST'])
def export_prd_file():
    """Export PRD content in a specific format."""
    data = request.json
    content = data.get('content', '')
    filename = data.get('filename', 'prd_export')
    target_format = data.get('format', 'md').lower().strip('.')
    
    if not content:
        return jsonify({'error': 'No content to export'}), 400
        
    try:
        if target_format in ['docx', 'doc']:
            # Generate DOCX
            doc = docx.Document()
            doc.add_heading('Product Requirements Document', 0)
            
            # Simple markdown parsing for the docx
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)
            
            byte_io = io.BytesIO()
            doc.save(byte_io)
            byte_io.seek(0)
            
            return send_file(
                byte_io,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"{filename}.docx"
            )
            
        elif target_format == 'pdf':
            # Fallback to DOCX or return error as PDF requires more setup
            return jsonify({'error': 'PDF export not yet implemented. Please use DOCX or MD.'}), 501
            
        else:
            # Default to text/markdown
            byte_io = io.BytesIO(content.encode('utf-8'))
            byte_io.seek(0)
            
            return send_file(
                byte_io,
                mimetype='text/markdown' if target_format == 'md' else 'text/plain',
                as_attachment=True,
                download_name=f"{filename}.{target_format}"
            )
            
    except Exception as e:
        logger.error(f"Export error: {e}")
        return jsonify({'error': f'Export failed: {str(e)}'}), 500

# Roadmap API Endpoints
@app.route('/api/roadmap', methods=['GET'])
@login_required
def list_roadmap_requests():
    return jsonify(get_roadmaps())

@app.route('/api/roadmap', methods=['POST'])
@login_required
def create_roadmap_request():
    try:
        # We handle multipart form for file uploads
        request_id = str(uuid.uuid4())
        
        # Basic Info
        data = {
            'id': request_id,
            'title': request.form.get('title'),
            'description': request.form.get('description'),
            'business_unit': request.form.get('business_unit'),
            'target_year': request.form.get('target_year'),
            'half_year': request.form.get('half_year'),
            'quarter': request.form.get('quarter'),
            'feature_type': request.form.get('feature_type'),
            'business_impact': int(request.form.get('business_impact', 1)),
            'created_at': datetime.now().isoformat(),
            'dependencies': json.loads(request.form.get('dependencies', '[]')),
            'attachments': []
        }
        
        # Handlee File Uploads
        request_upload_dir = os.path.join(ROADMAP_UPLOAD_DIR, request_id)
        os.makedirs(request_upload_dir, exist_ok=True)
        
        for file_key in ['prd_file', 'mockups_file']:
            if file_key in request.files:
                file = request.files[file_key]
                if file and file.filename:
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(request_upload_dir, filename)
                    file.save(filepath)
                    data['attachments'].append({
                        'type': 'prd' if file_key == 'prd_file' else 'mockup',
                        'filename': filename,
                        'path': filepath
                    })
        
        roadmaps = get_roadmaps()
        roadmaps.append(data)
        save_roadmaps(roadmaps)
        
        return jsonify({'success': True, 'id': request_id})
    except Exception as e:
        logger.error(f"Create roadmap error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/roadmap/<id>', methods=['GET'])
@login_required
def get_roadmap_request(id):
    roadmaps = get_roadmaps()
    request = next((r for r in roadmaps if r['id'] == id), None)
    if request:
        return jsonify(request)
    return jsonify({'error': 'Not found'}), 404

@app.route('/api/roadmap/<id>', methods=['PUT'])
def update_roadmap_request(id):
    try:
        roadmaps = get_roadmaps()
        idx = next((i for i, r in enumerate(roadmaps) if r['id'] == id), None)
        
        if idx is None:
            return jsonify({'error': 'Not found'}), 404
            
        # For simplicity, we expect JSON for updates (no new file uploads via PUT in this simple version)
        update_data = request.json
        roadmaps[idx].update(update_data)
        save_roadmaps(roadmaps)
        
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Update roadmap error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/roadmap/<id>', methods=['DELETE'])
@login_required
def delete_roadmap_request(id):
    try:
        roadmaps = get_roadmaps()
        roadmaps = [r for r in roadmaps if r['id'] != id]
        save_roadmaps(roadmaps)
        
        # Delete associated files
        request_upload_dir = os.path.join(ROADMAP_UPLOAD_DIR, id)
        if os.path.exists(request_upload_dir):
            shutil.rmtree(request_upload_dir)
            
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Delete roadmap error: {e}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    print("Starting AI PRD Management...")
    print(f"Upload folder: {app.config['UPLOAD_FOLDER']}")
    print(f"OpenAI available: {client is not None}")
    app.run(debug=True, host='127.0.0.1', port=5000)