import os
import shutil
import logging
import uuid
import json
from datetime import datetime
import io
import docx
from typing import Optional, List, Dict, Any
from contextlib import asynccontextmanager

from fastapi import FastAPI, Request, Form, UploadFile, File, Depends, HTTPException, status
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.security import OAuth2PasswordBearer
from starlette.middleware.sessions import SessionMiddleware
from starlette.middleware.cors import CORSMiddleware
from authlib.integrations.starlette_client import OAuth

from dependencies import (
    templates, get_users_data, save_users_data, 
    get_roadmaps_data, save_roadmaps_data,
    get_llm_client, UPLOAD_DIR, FORM_CONFIG_FILE,
    ROADMAP_FILE
)
from services.parser import parse_document, parse_prd_structure

# Logging
logger = logging.getLogger("aop_planner.main")

# App Setup
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    logger.info("Starting AOP Planner (FastAPI)...")
    yield
    # Shutdown
    logger.info("Shutting down...")

app = FastAPI(lifespan=lifespan, title="AOP Planner")

# Middleware
app.add_middleware(
    SessionMiddleware, 
    secret_key=os.getenv('SECRET_KEY', 'fw-aop-secret-key-2025'),
    https_only=False # Set to True in production
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Static Files
app.mount("/static", StaticFiles(directory="static"), name="static")
# Note: Original app might not have had a static folder explicit, templates might leverage inline or root. 
# Detailed check of file listing showed 'uploads' and 'templates'. I will mount uploads as well if needed, 
# but usually serving uploads directly is risky. Keeping it for compatibility if templates reference it.
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

# Auth Setup
oauth = OAuth()
oauth.register(
    name='google',
    client_id=os.getenv('GOOGLE_CLIENT_ID'),
    client_secret=os.getenv('GOOGLE_CLIENT_SECRET'),
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={'scope': 'openid email profile'},
)

# Helpers
def get_current_user(request: Request):
    user_data = request.session.get('user')
    if not user_data:
        return None
    return user_data

def login_required(request: Request):
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=302, headers={"Location": "/login"})
    return user

def admin_required(request: Request):
    user = login_required(request)
    if user.get('role') != 'admin':
        raise HTTPException(status_code=403, detail="Admin access required")
    return user

# --- Routes ---

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    user = get_current_user(request)
    if not user:
        return RedirectResponse("/login")
    return templates.TemplateResponse("home.html", {"request": request, "current_user": user})

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    if get_current_user(request):
        return RedirectResponse("/")
    return templates.TemplateResponse("login.html", {"request": request})

@app.get("/login/google")
async def login_google(request: Request):
    redirect_uri = request.url_for('auth_google_callback')
    return await oauth.google.authorize_redirect(request, redirect_uri)

@app.get("/login/google/callback", name="auth_google_callback")
async def auth_google_callback(request: Request):
    token = await oauth.google.authorize_access_token(request)
    user_info = token.get('userinfo')
    if not user_info:
        # fallback if not in token, fetch from endpoint
        # user_info = await oauth.google.parse_id_token(request, token) # simplified
        user_info = dict(token) # simplifying for now, assuming standard flow

    # Logic to find/create user
    users = get_users_data()
    email = user_info.get('email')
    user_data = next((u for u in users if u.get('email') == email), None)
    
    if not user_data:
        user_data = {
            "id": str(uuid.uuid4()),
            "username": email,
            "name": user_info.get('name'),
            "email": email,
            "role": "requester",
            "avatar": user_info.get('picture', f"https://ui-avatars.com/api/?name={user_info.get('name', 'User')}")
        }
        users.append(user_data)
        save_users_data(users)
    
    request.session['user'] = user_data
    return RedirectResponse("/")

@app.get("/logout")
async def logout(request: Request):
    request.session.pop('user', None)
    return RedirectResponse("/login")

# Simple API Login for testing
@app.post("/api/login")
async def api_login(request: Request):
    data = await request.json()
    username = data.get('username')
    password = data.get('password')
    
    users = get_users_data()
    user_data = next((u for u in users if u.get('username') == username and u.get('password') == password), None)
    
    if user_data:
        request.session['user'] = user_data
        return {"success": True}
    
    return JSONResponse({"success": False, "error": "Invalid credentials"}, status_code=401)

# Pages
@app.get("/prd", response_class=HTMLResponse)
async def prd_tool(request: Request):
    user = login_required(request)
    return templates.TemplateResponse("prd_management.html", {"request": request, "current_user": user})

@app.get("/roadmap", response_class=HTMLResponse)
async def roadmap_tool(request: Request):
    user = login_required(request)
    return templates.TemplateResponse("roadmap.html", {"request": request, "current_user": user})

@app.get("/admin", response_class=HTMLResponse)
async def admin_panel(request: Request):
    user = admin_required(request)
    return templates.TemplateResponse("admin.html", {"request": request, "current_user": user})

# API Endpoints
@app.get("/api/form-config")
async def get_form_config(request: Request):
    login_required(request)
    if FORM_CONFIG_FILE.exists():
        return json.loads(FORM_CONFIG_FILE.read_text())
    return []

@app.post("/api/form-config")
async def update_form_config(request: Request):
    admin_required(request)
    config = await request.json()
    FORM_CONFIG_FILE.write_text(json.dumps(config, indent=2))
    return {"success": True}

@app.post("/api/upload")
async def upload_prd_file(file: UploadFile = File(...)):
    # Note: Flask code had specific logic for parsing. Migrating it here.
    filename = file.filename
    if not filename:
         raise HTTPException(status_code=400, detail="No file")
    
    # Save file
    filepath = UPLOAD_DIR / filename
    with open(filepath, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
        
    try:
        # Parse
        parsed_data = parse_document(str(filepath), filename)
        structure = parse_prd_structure(parsed_data['content'])
        
        metadata = {
            'original_filename': filename,
            'uploaded_at': datetime.now().isoformat(),
            'file_size': len(content),
            'parsed_data': parsed_data,
            'structure': structure
        }
        
        # Save metadata
        metadata_file = UPLOAD_DIR / f"{filename}.meta.json"
        metadata_file.write_text(json.dumps(metadata, indent=2))
        
        return {
            'success': True,
            'filename': filename,
            'parsed_data': parsed_data,
            'structure': structure,
            'metadata': metadata
        }
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/api/list")
async def list_prds(request: Request):
    try:
        prds = []
        for item in os.listdir(UPLOAD_DIR):
            if item.endswith('.meta.json'):
                meta_path = UPLOAD_DIR / item
                try:
                    metadata = json.loads(meta_path.read_text())
                    prds.append({
                        'filename': metadata.get('original_filename', item.replace('.meta.json', '')),
                        'metadata': metadata
                    })
                except: 
                    pass
        return {'success': True, 'prds': prds, 'count': len(prds)}
    except Exception as e:
         return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/save")
async def save_prd(request: Request):
    """Save edited PRD."""
    login_required(request)
    data = await request.json()
    content = data.get('content', '')
    filename = data.get('filename')
    
    if not filename:
        filename = f'prd_edited_{datetime.now().strftime("%Y%m%d_%H%M%S")}.md'
    
    if not content:
        raise HTTPException(status_code=400, detail="No content to save")
    
    try:
        # Save the edited content
        save_path = UPLOAD_DIR / filename
        save_path.write_text(content, encoding='utf-8')
            
        # Parse for metadata consistency
        structure = parse_prd_structure(content)
        metadata = {
            'original_filename': filename,
            'uploaded_at': datetime.now().isoformat(),
            'file_size': len(content.encode('utf-8')),
            'parsed_data': {
                'content': content,
                'filename': filename
            },
            'structure': structure
        }
        
        # Save metadata to file
        metadata_file = UPLOAD_DIR / f"{filename}.meta.json"
        metadata_file.write_text(json.dumps(metadata, indent=2))
        
        return {
            'success': True,
            'filename': filename,
            'saved_at': datetime.now().isoformat(),
            'metadata': metadata
        }
    except Exception as e:
        logger.error(f"Save error: {e}")
        return JSONResponse({"error": f'Save failed: {str(e)}'}, status_code=500)

async def get_rice_analysis(content: str):
    """Extract RICE scores from PRD content using AI."""
    client = get_llm_client()
    if not client:
        return None
        
    try:
        response = client.chat.completions.create(
            model="Azure-GPT-5-chat",
            messages=[
                {"role": "system", "content": "You are a product management consultant. Analyze the provided PRD and extract RICE (Reach, Impact, Confidence, Effort) scores.\n\nReach: Number of users/customers affected (e.g., 500, 10000).\nImpact: Value to the user (0.25 to 3.0: 3=Massive, 2=High, 1=Medium, 0.5=Low, 0.25=Minimal).\nConfidence: Percentage certainty (50% to 100%).\nEffort: Person-months (e.g., 0.5, 3.0).\n\nReturn ONLY a JSON object with keys: reach, impact, confidence, effort, total_score, and verdict (a 1-sentence strategic summary)."},
                {"role": "user", "content": f"Analyze this PRD for RICE scores:\n\n{content}"}
            ],
            response_format={ "type": "json_object" }
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        logger.error(f"RICE extraction error: {e}")
        return None

@app.post("/api/analyze")
async def analyze_prd_endpoint(request: Request):
    """Analyze PRD for completeness and RICE prioritization."""
    login_required(request)
    data = await request.json()
    prd_content = data.get('content', '')
    
    if not prd_content:
        raise HTTPException(status_code=400, detail="No content provided")
    
    try:
        # 1. Structure Analysis
        structure = parse_prd_structure(prd_content)
        
        # 2. Completeness Check
        missing_sections = []
        for section, section_content in structure.items():
            if not section_content.strip():
                missing_sections.append(section.replace('_', ' ').title())
        
        total_sections = len(structure)
        filled_sections = sum(1 for section_content in structure.values() if section_content.strip())
        completeness_score = (filled_sections / total_sections) * 100
        
        # 3. RICE Analysis
        rice_data = await get_rice_analysis(prd_content)
        
        # 4. Final Payload
        return {
            'success': True,
            'analysis': {
                'completeness_score': round(completeness_score, 2),
                'missing_sections': missing_sections,
                'word_count': len(prd_content.split()),
                'section_counts': {k: len(v.split()) for k, v in structure.items()},
                'rice': rice_data,
                'recommendations': [
                    "Add clear success metrics" if not structure['success_metrics'] else "",
                    "Define user personas" if not structure['user_stories'] else "",
                    "Include timeline estimates" if not structure['timeline'] else ""
                ]
            }
        }
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        return JSONResponse({"error": f'Analysis failed: {str(e)}'}, status_code=500)

@app.post("/api/export")
async def export_prd_endpoint(request: Request):
    """Export PRD content in a specific format."""
    data = await request.json()
    content = data.get('content', '')
    filename = data.get('filename', 'prd_export')
    target_format = data.get('format', 'md').lower().strip('.')
    
    if not content:
        raise HTTPException(status_code=400, detail="No content to export")
        
    try:
        if target_format in ['docx', 'doc']:
            # Generate DOCX
            doc = docx.Document()
            doc.add_heading('Product Requirements Document', 0)
            
            # Simple markdown parsing for the docx
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)
            
            byte_io = io.BytesIO()
            doc.save(byte_io)
            byte_io.seek(0)
            
            headers = {
                'Content-Disposition': f'attachment; filename="{filename}.docx"'
            }
            from fastapi.responses import StreamingResponse
            return StreamingResponse(byte_io, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=headers)
        
        # Default to MD
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}.md"'
        }
        return StreamingResponse(io.BytesIO(content.encode('utf-8')), media_type='text/markdown', headers=headers)
    except Exception as e:
        logger.error(f"Export error: {e}")
        return JSONResponse({"error": f'Export failed: {str(e)}'}, status_code=500)

@app.post("/api/admin/load-sample")
async def load_sample_data(request: Request):
    admin_required(request)
    try:
        # Logic to append to roadmaps.json
        from dependencies import get_roadmaps_data, save_roadmaps_data
        
        sample_raw = [
            {"name": "Forms 2.0: Richer Fields, Seamless Migration, Unified Experience", "desc": "Introduce Rich Text Editor, Attachment, and Formula fields in all forms...", "bu": "EX", "type": "Major Feature", "priority": "P2", "q": "Q2"},
            {"name": "Automated Change Risk scoring and Assessment", "desc": "Introduce a native, scalable risk assessment framework within Change Management...", "bu": "EX", "type": "Major Feature", "priority": "P1", "q": "Q1"},
            {"name": "Service Health Dashboard powered by D42-ITOM Integration", "desc": "A more powerful and feature-rich Service Health Dashboard powered by Device42 and ITOM", "bu": "EX", "type": "Major Feature", "priority": "P1", "q": "Q1"},
            {"name": "New FD Omni for existing customers", "desc": "Migration readiness for customers from FD(older accounts), Classic Omni & CSS Omni...", "bu": "CX", "type": "Major Feature", "priority": "P1", "q": "Q2"},
            {"name": "Proactive messaging for Omnichannel Freshdesk", "desc": "Support for contact segmentation and proactive campaigns across channels...", "bu": "CX", "type": "Major Feature", "priority": "P2", "q": "Q3"},
            {"name": "Enhanced portal", "desc": "AI powered search on the portal, AI prompt portal customization, OOTB themes...", "bu": "CX", "type": "Major Feature", "priority": "P2", "q": "Q2"},
            {"name": "Enhanced Knowledge base", "desc": "AI powered KB insights and article generation, AI powered translation...", "bu": "CX", "type": "Major Feature", "priority": "P2", "q": "Q3"},
            {"name": "Multi-departments", "desc": "Improved RBAC and isolation for tickets, contacts, analytics and admin setup...", "bu": "CX", "type": "Major Feature", "priority": "P2", "q": "Q1"}
        ]

        roadmaps = get_roadmaps_data()
        
        new_items = []
        for s in sample_raw:
            bu_map = {"EX": "EX BU", "AI": "AI BU", "CX": "CX BU", "CE": "CE BU"}
            impact_map = {"P1": 5, "P2": 4, "P3": 3}
            
            item = {
                "id": str(uuid.uuid4()),
                "name": s['name'],
                "description": s['desc'],
                "bu": bu_map.get(s['bu'], "AI BU"),
                "type": "Hero Big Rock" if "Major" in s['type'] else "Big Rock",
                "impact": impact_map.get(s['priority'], 3),
                "confidence": 80,
                "effort": 3,
                "quarter": s['q'],
                "status": "Draft",
                "created_at": datetime.now().isoformat()
            }
            new_items.append(item)
            
        roadmaps.extend(new_items)
        save_roadmaps_data(roadmaps)
        
        return {"success": True, "count": len(new_items)}
    except Exception as e:
        logger.error(f"Load sample error: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)
async def delete_prd(filename: str, request: Request):
    login_required(request)
    # Simple security check to prevent directory traversal
    if ".." in filename or "/" in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
        
    filepath = UPLOAD_DIR / filename
    meta_path = UPLOAD_DIR / f"{filename}.meta.json"
    
    if filepath.exists():
        os.remove(filepath)
    if meta_path.exists():
        os.remove(meta_path)
        
    return {"success": True}

ROADMAP_ATTACHMENTS_DIR = UPLOAD_DIR / "roadmap_attachments"
ROADMAP_ATTACHMENTS_DIR.mkdir(exist_ok=True)

@app.get("/api/roadmap")
async def list_roadmap_requests(request: Request):
    login_required(request)
    return get_roadmaps_data()

@app.post("/api/roadmap")
async def create_roadmap_request(request: Request):
    login_required(request)
    try:
        # Note: FastAPI handles multipart/form-data via Form and UploadFile
        # But for dynamic fields it might be easier to use the request object
        form_data = await request.form()
        request_id = str(uuid.uuid4())
        
        # Basic Info
        data = {
            'id': request_id,
            'title': form_data.get('title'),
            'description': form_data.get('description'),
            'business_unit': form_data.get('business_unit'),
            'target_year': form_data.get('target_year'),
            'half_year': form_data.get('half_year'),
            'quarter': form_data.get('quarter'),
            'feature_type': form_data.get('feature_type'),
            'business_impact': int(form_data.get('business_impact', 1)),
            'created_at': datetime.now().isoformat(),
            'dependencies': json.loads(form_data.get('dependencies', '[]')),
            'attachments': []
        }
        
        # Handle File Uploads
        request_upload_dir = ROADMAP_ATTACHMENTS_DIR / request_id
        request_upload_dir.mkdir(exist_ok=True)
        
        for file_key in ['prd_file', 'mockups_file']:
            if file_key in form_data:
                file = form_data[file_key]
                if hasattr(file, 'filename') and file.filename:
                    filename = file.filename
                    filepath = request_upload_dir / filename
                    with open(filepath, "wb") as buffer:
                        buffer.write(await file.read())
                    data['attachments'].append({
                        'type': 'prd' if file_key == 'prd_file' else 'mockup',
                        'filename': filename,
                        'path': str(filepath)
                    })
        
        roadmaps = get_roadmaps_data()
        roadmaps.append(data)
        save_roadmaps_data(roadmaps)
        
        return {'success': True, 'id': request_id}
    except Exception as e:
        logger.error(f"Create roadmap error: {e}")
        return JSONResponse({'error': str(e)}, status_code=500)

@app.get("/api/roadmap/{id}")
async def get_roadmap_request(id: str, request: Request):
    login_required(request)
    roadmaps = get_roadmaps_data()
    req = next((r for r in roadmaps if r['id'] == id), None)
    if req:
        return req
    raise HTTPException(status_code=404, detail="Not found")

@app.put("/api/roadmap/{id}")
async def update_roadmap_request(id: str, request: Request):
    login_required(request)
    try:
        roadmaps = get_roadmaps_data()
        idx = next((i for i, r in enumerate(roadmaps) if r['id'] == id), None)
        
        if idx is None:
            raise HTTPException(status_code=404, detail="Not found")
            
        update_data = await request.json()
        roadmaps[idx].update(update_data)
        save_roadmaps_data(roadmaps)
        
        return {'success': True}
    except Exception as e:
        logger.error(f"Update roadmap error: {e}")
        return JSONResponse({'error': str(e)}, status_code=500)

@app.delete("/api/roadmap/{id}")
async def delete_roadmap_request(id: str, request: Request):
    login_required(request)
    try:
        roadmaps = get_roadmaps_data()
        initial_len = len(roadmaps)
        roadmaps = [r for r in roadmaps if r['id'] != id]
        
        if len(roadmaps) == initial_len:
             raise HTTPException(status_code=404, detail="Not found")
             
        save_roadmaps_data(roadmaps)
        
        # Delete associated files
        request_upload_dir = ROADMAP_ATTACHMENTS_DIR / id
        if request_upload_dir.exists():
            shutil.rmtree(request_upload_dir)
            
        return {'success': True}
    except Exception as e:
        logger.error(f"Delete roadmap error: {e}")
        return JSONResponse({'error': str(e)}, status_code=500)

from routers import workflow
app.include_router(workflow.router)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
